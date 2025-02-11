from docx import Document
from docx.enum.text import WD_BREAK
import re

def split_questions_in_place(docx_file):
    # Load the existing document
    doc = Document(docx_file)
    
    # Regular expression to match question numbers like "1)", "2)", etc.
    question_pattern = re.compile(r'^\d+\)')

    # Iterate through paragraphs in reverse order
    for i in range(len(doc.paragraphs) - 1, -1, -1):
        para = doc.paragraphs[i]
        
        # Check if the paragraph starts with a question number
        if question_pattern.match(para.text.strip()):
            if i != 0:  # Avoid adding a page break before the first question
                # Insert a page break before this paragraph
                doc.paragraphs[i].insert_paragraph_before().add_run().add_break(WD_BREAK.PAGE)

    # Save the changes to the original document
    doc.save(docx_file)


